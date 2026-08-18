"""DOCX parser. Extracts every table via python-docx, then runs each grid
through header/region detection. Import-guarded.
"""
from __future__ import annotations

import io
from typing import List

from ..extraction.table_detection import tables_from_grid
from ..models.documents import (ExtractionResult, FinancialDocument, FinancialTable,
                                 ParseWarning)
from ..models.enums import FileFormat, Severity, WarningCode

try:
    import docx
    _DOCX_OK = True
except Exception:  # pragma: no cover
    _DOCX_OK = False


class DocxParser:
    formats = [FileFormat.DOCX]

    def parse(self, filename: str, data: bytes, pages=None) -> ExtractionResult:
        """`pages` selects 1-based *table* indices (Word has no fixed pages)."""
        if not _DOCX_OK:
            return _lib_missing(filename)
        wanted = set(int(p) for p in pages) if pages else None
        warnings: List[ParseWarning] = []
        document = docx.Document(io.BytesIO(data))
        tables: List[FinancialTable] = []

        for t_idx, table in enumerate(document.tables):
            if wanted is not None and (t_idx + 1) not in wanted:
                continue
            grid = [[cell.text for cell in row.cells] for row in table.rows]
            found = tables_from_grid(grid, filename, f"table_{t_idx + 1}")
            tables.extend(found)

        if not tables:
            warnings.append(ParseWarning(WarningCode.NO_TABLES_FOUND,
                                         f"No tables found in {filename}.", Severity.WARNING))
        doc = FinancialDocument(filename, FileFormat.DOCX, tables,
                                metadata={"docx_tables": len(document.tables)})
        return ExtractionResult(doc, warnings, {"tables": len(tables), "rows": doc.record_count})


def _lib_missing(filename: str) -> ExtractionResult:
    w = ParseWarning(WarningCode.PARSER_LIBRARY_MISSING,
                     "python-docx is not installed in this code env; "
                     f"{filename} was skipped. Add 'python-docx' to requirements.",
                     Severity.ERROR)
    return ExtractionResult(FinancialDocument(filename, FileFormat.DOCX, []), [w], {"rows": 0})
