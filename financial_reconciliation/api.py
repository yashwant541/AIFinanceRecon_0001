"""Public functional API — each pipeline step as a standalone function.

    Step 1  extract_tables / list_sheets      -> tables with detected headers
    Step 2  profile_side / suggest_mapping     -> auto key/value structure
            to_key_value_records               -> (keys) -> (values) per side
    Step 3  reconcile                          -> line/key match + variance
    Step 4  export_workbook                    -> downloadable .xlsx

These wrap the object-oriented engine so other apps can call any step in
isolation without touching the web layer.
"""
from __future__ import annotations

from typing import Dict, List, Optional, Set

from .config.settings import EngineConfig
from .extraction.profiling import ColumnProfile, profile_column
from .mapping.auto_map import suggest_mapping as _suggest
from .models.documents import (ExtractionResult, FinancialDocument, FinancialRecord,
                               FinancialTable, UploadedDocument)
from .models.enums import FileFormat, Side
from .models.results import CanonicalRecord, ReconciliationResult
from .normalization.normalizer import Normalizer
from .ontology import Ontology
from .parsers.base import default_factory
from .parsers.excel_parser import list_sheets as _list_sheets
from .pipeline import ReconciliationPipeline

_FACTORY = default_factory()
_PIPELINE = ReconciliationPipeline(_FACTORY)


# ---- Step 1: extraction --------------------------------------------------
def list_sheets(filename: str, data: bytes) -> List[str]:
    """Sheet names for an Excel workbook (for a sheet picker)."""
    return _list_sheets(filename, data)


def inspect_file(filename: str, data: bytes) -> Dict:
    """Describe a file so the user can choose what to read.

    Returns {kind, sheets[], units, count} where `units` names what `count`
    counts: 'sheets' for Excel, 'pages' for PDF, 'tables' for DOCX.
    """
    low = filename.lower()
    if low.endswith((".xlsx", ".xlsm", ".xls")):
        try:
            sheets = _list_sheets(filename, data)
        except Exception:  # noqa: BLE001
            sheets = []
        return {"kind": "excel", "sheets": sheets, "units": "sheets",
                "count": len(sheets)}
    if low.endswith(".pdf"):
        n = 0
        try:
            import io as _io
            import pdfplumber
            with pdfplumber.open(_io.BytesIO(data)) as pdf:
                n = len(pdf.pages)
        except Exception:  # noqa: BLE001
            n = 0
        return {"kind": "pdf", "sheets": [], "units": "pages", "count": n}
    if low.endswith(".docx"):
        n = 0
        try:
            import io as _io
            import docx as _docx
            n = len(_docx.Document(_io.BytesIO(data)).tables)
        except Exception:  # noqa: BLE001
            n = 0
        return {"kind": "docx", "sheets": [], "units": "tables", "count": n}
    return {"kind": "other", "sheets": [], "units": "", "count": 0}


def parse_ranges(spec: str, maximum: int) -> Optional[List[int]]:
    """'1-3,7' -> [1,2,3,7]. Empty/'all' -> None (meaning everything)."""
    if spec is None:
        return None
    s = str(spec).strip().lower()
    if s in ("", "all", "*"):
        return None
    out: List[int] = []
    for part in s.replace(" ", "").split(","):
        if not part:
            continue
        if "-" in part:
            a, _, b = part.partition("-")
            try:
                lo, hi = int(a), int(b)
            except ValueError:
                continue
            out.extend(range(min(lo, hi), max(lo, hi) + 1))
        else:
            try:
                out.append(int(part))
            except ValueError:
                continue
    out = sorted({p for p in out if 1 <= p <= (maximum or p)})
    return out or None


def list_engines() -> List[str]:
    """PDF extraction engines available in this environment."""
    from .extraction.engines import available_engines
    return available_engines()


def extract_tables(filename: str, data: bytes,
                   sheets: Optional[List[str]] = None,
                   pages: Optional[List[int]] = None,
                   first_table_only: bool = False,
                   engine: str = "native") -> ExtractionResult:
    """Parse one file into tables with detected headers.

    `sheets` filters Excel worksheets; `pages` filters PDF pages (or DOCX table
    indices); `first_table_only` keeps just the first table per Excel sheet;
    `engine` selects the PDF backend (native / camelot / camelot+native).
    """
    if filename.lower().endswith(".json"):
        return load_processed(data)
    parser = _FACTORY.for_filename(filename)
    kind = parser.__class__.__name__
    if kind == "ExcelParser":
        return parser.parse(filename, data, sheets=sheets,
                            first_table_only=first_table_only)  # type: ignore[call-arg]
    if kind == "PdfParser" and engine and engine != "native":
        from .extraction.engines import extract_with_engine
        from .models.documents import FinancialDocument
        from .models.enums import FileFormat as _FF

        def _native(fn, dt, pg):
            return parser.parse(fn, dt, pages=pg)  # type: ignore[call-arg]

        tables = extract_with_engine(filename, data, engine, _native, pages)
        doc = FinancialDocument(filename, _FF.PDF, tables,
                                metadata={"engine": engine})
        return ExtractionResult(doc, [], {"tables": len(tables),
                                          "rows": doc.record_count})
    if kind in ("PdfParser", "DocxParser"):
        return parser.parse(filename, data, pages=pages)  # type: ignore[call-arg]
    return parser.parse(filename, data)


def extract_many(docs: List[UploadedDocument]) -> List[ExtractionResult]:
    out: List[ExtractionResult] = []
    for d in docs:
        if d.filename.lower().endswith(".json"):
            out.append(load_processed(d.data))
        else:
            out.append(_FACTORY.for_filename(d.filename).parse(d.filename, d.data))
    return out


def table_id(source_file: str, table_name: str) -> str:
    return f"{source_file}||{table_name}"


def qc_table(table) -> Dict:
    """Re-derive subtotal chains to flag whether a table's numbers foot."""
    from .extraction.qc import reconcile_subtotals
    rows = [r.values for r in table.records]
    return reconcile_subtotals(table.columns, rows)


# ---- Ontology / synonyms -------------------------------------------------
def load_ontology(filename: str, data: bytes) -> Ontology:
    """Load a user-supplied synonym list (json/csv/txt)."""
    return Ontology.from_file(filename, data)


# ---- Reshape wide -> long (tidy) ----------------------------------------
def melt_extraction(extraction: ExtractionResult, period_name: str = "Period",
                    value_name: str = "Value") -> ExtractionResult:
    """Return a copy where every wide table is unpivoted to [dims..., Value]."""
    from .extraction.reshape import melt
    tables = []
    for t in extraction.document.tables:
        cols, rows = melt([c for c in t.columns], [r.values for r in t.records],
                          period_name, value_name)
        recs = [FinancialRecord(values=row, source_file=t.source_file,
                                table_name=t.name, row_index=i)
                for i, row in enumerate(rows)]
        tables.append(FinancialTable(t.name, cols, recs, t.source_file))
    doc = FinancialDocument(extraction.document.filename,
                            extraction.document.file_format, tables,
                            dict(extraction.document.metadata))
    return ExtractionResult(doc, list(extraction.warnings),
                            {"tables": len(tables), "rows": doc.record_count})


# ---- Processed-document round-trip (reference library) -------------------
def export_processed(extractions: List[ExtractionResult], source_name: str,
                     long: bool = True) -> Dict:
    """Serialize extracted tables to a reusable 'processed document' dict.

    long=True emits the tidy [dims..., Value] schema used by the reference
    library; long=False keeps the wide layout.
    """
    tables = []
    for e in extractions:
        src = melt_extraction(e) if long else e
        for t in src.document.tables:
            tables.append({"name": t.name, "columns": t.columns,
                           "rows": [r.values for r in t.records]})
    return {"source": source_name, "format": "finrecon.processed.v1",
            "layout": "long" if long else "wide", "tables": tables}


def load_processed(data) -> ExtractionResult:
    """Rebuild an ExtractionResult from a processed-document dict or json bytes."""
    import json
    obj = data if isinstance(data, dict) else json.loads(
        data.decode("utf-8-sig", errors="replace") if isinstance(data, (bytes, bytearray))
        else data)
    source = obj.get("source", "processed")
    tables = []
    for t in obj.get("tables", []):
        cols = t.get("columns") or (list(t["rows"][0].keys()) if t.get("rows") else [])
        recs = [FinancialRecord(values=row, source_file=source,
                                table_name=t.get("name", "table"), row_index=i)
                for i, row in enumerate(t.get("rows", []))]
        tables.append(FinancialTable(t.get("name", "table"), cols, recs, source))
    doc = FinancialDocument(source, FileFormat.CSV, tables, metadata={"processed": True})
    return ExtractionResult(doc, [], {"tables": len(tables), "rows": doc.record_count})


# ---- Step 2: profiling, mapping, key/value pairs -------------------------
def profile_side(extractions: List[ExtractionResult]) -> List[ColumnProfile]:
    """One profile per distinct column across a side's tables."""
    columns: List[str] = []
    for e in extractions:
        for c in e.all_columns():
            if c not in columns:
                columns.append(c)
    records = [r for e in extractions for r in e.all_records()]
    return [profile_column(c, [r.values.get(c) for r in records]) for c in columns]


def suggest_mapping(left: List[ExtractionResult],
                    right: List[ExtractionResult]) -> List[Dict]:
    """Auto-suggest canonical fields aligning both sides' columns."""
    return _suggest(profile_side(left), profile_side(right))


def to_key_value_records(extractions: List[ExtractionResult], config: EngineConfig,
                         side: Side, ontology: Optional[Ontology] = None
                         ) -> List[CanonicalRecord]:
    """Step 2 made explicit: rows -> (key columns) -> (value columns), aggregated."""
    config.validate()
    records = [r for e in extractions for r in e.all_records()]
    return Normalizer(config, ontology=ontology).canonicalize(records, side)


# ---- Step 3: reconcile ---------------------------------------------------
def reconcile(left: List[ExtractionResult], right: List[ExtractionResult],
              config: EngineConfig,
              included_left: Optional[Set[str]] = None,
              included_right: Optional[Set[str]] = None,
              ontology: Optional[Ontology] = None,
              melt: bool = False) -> ReconciliationResult:
    if melt:
        def _long(e):  # already tidy (processed reference or a Value column)
            return bool(e.document.metadata.get("processed")) or \
                any("Value" in t.columns for t in e.document.tables)
        left = [e if _long(e) else melt_extraction(e) for e in left]
        right = [e if _long(e) else melt_extraction(e) for e in right]
    return _PIPELINE.reconcile(left, right, config, included_left, included_right,
                               ontology=ontology)


# ---- Step 4: export ------------------------------------------------------
def export_workbook(result: ReconciliationResult) -> bytes:
    from .exporters.workbook import WorkbookExporter
    return WorkbookExporter().export(result)


def result_to_dict(result: ReconciliationResult, cap: int = 250) -> Dict:
    from .exporters.serialize import result_to_dict as _r2d
    return _r2d(result, cap)
